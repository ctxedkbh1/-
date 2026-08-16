from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.exporters.style_parser import parse_format_note

ALIGN_MAP = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
             "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}


def _add_page_number_footer(section, font_cn="宋体"):
    from docx.oxml import OxmlElement
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._element.append(fld1)
    run._element.append(instr)
    run._element.append(fld2)


class DocxExporter:
    def export(self, doc, path):
        fmt = doc.meta.get("format_dict") or parse_format_note(doc.meta.get("format_note", ""))
        font = fmt.get("font", "宋体")
        latin = fmt.get("latin_font", "Times New Roman")
        size = fmt.get("size", 12.0)
        spacing = fmt.get("line_spacing", 1.5)
        indent = fmt.get("first_indent_chars", 2)
        align = ALIGN_MAP.get(fmt.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)

        d = DocxDocument()
        section = d.sections[0]
        if fmt.get("page") == "A4":
            section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.top_margin = Cm(float(fmt.get("margin_top", 2.54)))
        section.bottom_margin = Cm(float(fmt.get("margin_bottom", 2.54)))
        section.left_margin = Cm(float(fmt.get("margin_left", 3.17)))
        section.right_margin = Cm(float(fmt.get("margin_right", 3.17)))

        normal = d.styles["Normal"]
        normal.font.name = latin
        normal.font.size = Pt(size)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), font)

        def set_cn(run, font_cn=None, rsize=None, bold=False):
            run.font.name = latin
            run.font.size = Pt(rsize if rsize is not None else size)
            run.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn or font)

        def para(text, rsize=None, bold=False, p_align=None, p_indent=True,
                 p_spacing=None, hanging=False, font_cn=None):
            p = d.add_paragraph()
            p.paragraph_format.line_spacing = p_spacing or spacing
            p.alignment = p_align if p_align is not None else align
            if p_indent:
                p.paragraph_format.first_line_indent = Pt((rsize or size) * indent)
            if hanging:
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(-1.0)
            set_cn(p.add_run(text), font_cn=font_cn, rsize=rsize, bold=bold)
            return p

        title = doc.meta.get("title") or "论文"
        para(title, rsize=fmt.get("title_size", 22), bold=bool(fmt.get("title_bold", True)),
             p_align=ALIGN_MAP.get(fmt.get("title_align", "center"), WD_ALIGN_PARAGRAPH.CENTER),
             p_indent=False, font_cn=fmt.get("title_font", "黑体"))
        info_line = "　".join(x for x in (
            doc.meta.get("name") and f"姓名：{doc.meta['name']}",
            doc.meta.get("school") and f"学校：{doc.meta['school']}",
            doc.meta.get("college") and f"学院：{doc.meta['college']}",
            doc.meta.get("major") and f"专业：{doc.meta['major']}",
            doc.meta.get("class_name") and f"班级：{doc.meta['class_name']}",
            doc.meta.get("student_id") and f"学号：{doc.meta['student_id']}",
            doc.meta.get("teacher") and f"指导老师：{doc.meta['teacher']}",
        ) if x)
        if info_line:
            para(info_line, p_align=WD_ALIGN_PARAGRAPH.CENTER, p_indent=False)

        if doc.abstract:
            p = d.add_paragraph()
            p.paragraph_format.line_spacing = spacing
            p.alignment = align
            set_cn(p.add_run("摘要："), font_cn=fmt.get("h1_font", "黑体"),
                   rsize=size, bold=True)
            set_cn(p.add_run(doc.abstract), rsize=size)
        if doc.keywords:
            p = d.add_paragraph()
            p.paragraph_format.line_spacing = spacing
            p.alignment = align
            set_cn(p.add_run("关键词："), font_cn=fmt.get("h1_font", "黑体"),
                   rsize=size, bold=True)
            set_cn(p.add_run("；".join(doc.keywords)), rsize=size)

        for block in doc.blocks:
            kind = block.get("kind")
            text = block.get("text", "")
            if kind == "heading1":
                para(text, rsize=fmt.get("h1_size", 15), bold=bool(fmt.get("h1_bold", True)),
                     p_indent=False, font_cn=fmt.get("h1_font", "黑体"))
            elif kind == "heading2":
                para(text, rsize=fmt.get("h2_size", 14), bold=bool(fmt.get("h2_bold", True)),
                     p_indent=False, font_cn=fmt.get("h2_font", "黑体"))
            elif kind == "heading3":
                para(text, rsize=fmt.get("h3_size", 12), bold=bool(fmt.get("h3_bold", True)),
                     p_indent=False, font_cn=fmt.get("h3_font", "黑体"))
            else:
                para(text, rsize=size, p_indent=True)

        if doc.references:
            d.add_page_break()
            para("参考文献", rsize=fmt.get("h1_size", 15), bold=True, p_indent=False,
                 font_cn=fmt.get("h1_font", "黑体"))
            for ref in doc.references:
                para(ref, rsize=10.5, p_indent=False, hanging=True)

        if doc.data_sources:
            d.add_page_break()
            para("数据来源", rsize=fmt.get("h1_size", 15), bold=True, p_indent=False,
                 font_cn=fmt.get("h1_font", "黑体"))
            for i, s in enumerate(doc.data_sources, 1):
                name = s.get("name", "")
                url = s.get("url", "")
                verified = s.get("verified", True)
                tag = "" if verified else "【待核实】"
                line = f"[{i}] {name}{tag}"
                if url:
                    line += f"　{url}"
                para(line, rsize=10.5, p_indent=False, hanging=True)

        if fmt.get("page_number") == "bottom_center":
            _add_page_number_footer(section)
        if fmt.get("header"):
            section.header.paragraphs[0].text = fmt.get("header")
            for run in section.header.paragraphs[0].runs:
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        d.save(path)

# 版本: v2.0.0 (2026-08-16) 更新: 高级模式工作台
